#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_DIR = Path(__file__).resolve().parent.parent
OUTPUT = PROJECT_DIR / "Instrucciones_Convertidor_OpenVAS.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F6FA")
    paragraph._p.get_or_add_pPr().append(shading)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)


def add_step(doc: Document, title: str, body: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)
    if body:
        run = paragraph.add_run(f" {body}")
        run.font.name = "Arial"
        run.font.size = Pt(10)


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(31, 78, 121)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Guia de uso - Convertidor OpenVAS XML a XLSX")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Exportacion manual desde Greenbone Web y conversion local a Excel")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.italic = True

    doc.add_heading("Objetivo", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Convertir reportes exportados desde Greenbone/OpenVAS en formato XML o Anonymous XML a un archivo "
        "Excel con hojas organizadas para analisis tecnico y remediacion."
    )

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(5.2)
    set_cell_shading(table.rows[0].cells[0], "1F4E79")
    set_cell_shading(table.rows[0].cells[1], "1F4E79")
    set_cell_text(table.rows[0].cells[0], "Elemento", True, "FFFFFF")
    set_cell_text(table.rows[0].cells[1], "Ubicacion / uso", True, "FFFFFF")
    rows = [
        ("Lanzador principal", "convertir_xml_a_xlsx.bat"),
        ("XML de entrada", "entradas"),
        ("Excel generado", "salidas"),
        ("Script base", "app\\openvas_report_to_xlsx.py"),
        ("Ejemplos", "ejemplos"),
    ]
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)

    doc.add_heading("Uso rapido", level=1)
    add_step(doc, "Exporta el reporte desde Greenbone.", "En la web abre el reporte y selecciona formato XML o Anonymous XML.")
    add_step(doc, "Guarda el archivo XML.", "Recomendado: copia el XML exportado a la carpeta entradas.")
    add_step(doc, "Ejecuta el BAT.", "Haz doble clic en convertir_xml_a_xlsx.bat.")
    add_step(doc, "Selecciona el XML.", "El asistente solo muestra reportes XML ubicados en la carpeta entradas.")
    add_step(doc, "Escribe el nombre del XLSX.", "No necesitas escribir la extension .xlsx; el asistente la agrega automaticamente.")
    add_step(doc, "Presiona Enter.", "Despues de elegir el XML y el nombre de salida, la conversion inicia sin pedir confirmacion adicional.")

    doc.add_heading("Comando manual equivalente", level=1)
    add_code_block(
        doc,
        '.\\.venv\\Scripts\\python.exe .\\app\\openvas_report_to_xlsx.py --xml .\\entradas\\reporte.xml --output .\\salidas\\reporte.xlsx',
    )

    doc.add_heading("Contenido del Excel", level=1)
    for item in [
        "Indice: enlaces internos a cada hoja y metadatos basicos del reporte.",
        "Resumen: datos generales, conteos y grafico por severidad.",
        "Vulnerabilidades: hallazgos detallados con host, puerto, severidad, QoD, CVEs y solucion.",
        "Por_Host: vista operativa para revisar hallazgos por activo.",
        "Hosts: resumen por activo afectado.",
        "Vuln_Detalle: vulnerabilidades deduplicadas con hosts afectados, impacto y solucion.",
        "CVEs: agrupacion por identificador CVE.",
        "Remediacion: lista priorizada de hallazgos con severidad media o superior.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Notas importantes", level=1)
    for item in [
        "En esta VM Greenbone OS 24.10.9 no se encontro GMP/API expuesto por TCP.",
        "El puerto 9390 no esta escuchando y el menu de servicios no muestra GMP.",
        "Por ese motivo el flujo recomendado es exportar XML desde la web y convertirlo localmente.",
        "No modifiques archivos internos del appliance; esto evita romper soporte, actualizaciones o firmas de formatos.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Solucion de problemas", level=1)
    troubleshooting = doc.add_table(rows=1, cols=2)
    troubleshooting.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_shading(troubleshooting.rows[0].cells[0], "1F4E79")
    set_cell_shading(troubleshooting.rows[0].cells[1], "1F4E79")
    set_cell_text(troubleshooting.rows[0].cells[0], "Problema", True, "FFFFFF")
    set_cell_text(troubleshooting.rows[0].cells[1], "Accion recomendada", True, "FFFFFF")
    rows = [
        ("No aparecen XML en entradas", "Copia el XML exportado desde la web a la carpeta entradas y vuelve a ejecutar el BAT."),
        ("El XLSX no se genera", "Verifica que el archivo sea XML completo de reporte, no HTML ni CSV."),
        ("Excel muestra reparar contenido", "Genera nuevamente el XLSX con esta version actualizada del convertidor."),
        ("El archivo queda bloqueado", "Cierra Excel y vuelve a ejecutar la conversion."),
        ("Faltan hallazgos", "Exporta el reporte completo desde Greenbone y evita filtros de filas en la vista web."),
    ]
    for key, value in rows:
        cells = troubleshooting.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Exportar_Informe - Convertidor OpenVAS XML a XLSX").font.size = Pt(8)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
